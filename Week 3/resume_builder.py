import os, json, time
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from pydantic import BaseModel, Field, ConfigDict


load_dotenv()
my_api_key=os.getenv("GOOGLE_API_KEY")
MODEL = "gemini-2.5-flash" 

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=OpenAI(
    api_key=my_api_key,
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
)
role="user"


# take resume in pdf or word
# have hr give you a list of things like skill, experience, projects
# extract these from resume 
# match against the hr list
# generate a percentage of matching or not

job_description="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""

# 1. take resume in pdf or word
from pypdf import PdfReader
from docx import Document


def parse_pdf(pdf_path: str) -> dict:
    """
    Parse a PDF and return its metadata and extracted text.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        dict: {
            "metadata": {...},
            "num_pages": int,
            "pages": [
                {"page": 1, "text": "..."},
                ...
            ],
            "full_text": "..."
        }
    """
    reader = PdfReader(pdf_path)

    # Extract metadata
    metadata = {}
    if reader.metadata:
        metadata = {
            key.lstrip("/"): str(value)
            for key, value in reader.metadata.items()
        }

    pages = []
    full_text = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append({
            "page": i + 1,
            "text": text
        })
        full_text.append(text)

    return {
        "metadata": metadata,
        "num_pages": len(reader.pages),
        "pages": pages,
        "full_text": "\n".join(full_text)
    }

def parse_docx(file_path: str):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def llm_call(messages):
    response_format={
        "type": "json_object"
    }
    
    response=client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format=response_format
    )

    return response.choices[0].message.content


# Change ALL model configs from extra='forbid' to extra='ignore'
# Apply to every model class

class MatchResult(BaseModel):
    score: float
    details: dict

class Project(BaseModel):
    model_config = ConfigDict(extra='ignore')   # was 'forbid' — crashes on unknown fields
    project_type: str | None = None
    summary: str | None = None

class Experience(BaseModel):
    model_config = ConfigDict(extra='ignore')   # add this — was missing entirely
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)

class Resume(BaseModel):
    model_config = ConfigDict(extra='ignore')   # add this — was missing entirely
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    experience: list[Experience] = Field(default_factory=list)
    skills: list[str] = Field(default_factory=list)
    projects: list[Project] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)

class JobDes(BaseModel):
    model_config = ConfigDict(extra='ignore')   # add this
    role: str
    minimum_experience: float | None = None
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    responsibility: str | None = None 



def parse_resume(result):
    schema = Resume.model_json_schema()   # was computed but never used — now we pass it

    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume and return JSON matching this schema EXACTLY:
    {json.dumps(schema, indent=2)}

    Rules:
    - Extract only information explicitly present in the resume. Never invent or guess.
    - Return ONLY the fields defined in the schema. No extra fields.
    - For projects, return ONLY these two fields:
        - project_type (string): one of: GenAI Project, AI/ML Project, FullStack Project,
        Python Backend, Backend with MongoDB, Frontend Project, Web Development Project,
        Data Science Project. Use null if unclear.
        - summary (string): 1-2 sentence description. Use null if not present.
    - Use null for missing scalar values, empty lists for missing list values.
    - Return valid JSON only. No markdown, no explanation.
    """

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user",   "content": f"Resume text:\n{result['full_text']}"},
    ]

    text = llm_call(messages)
    raw = json.loads(text)
    return Resume(**raw)

def parse_job(job_description):
    schema=JobDes.model_json_schema()
    
    jd_system_prompt = f"""
    You are a job description information extraction system.

    Your task is to extract structured information from the provided job description.

    Follow these rules:
    1. Extract information only from the job description.
    2. Do not infer, assume, or invent information that is not explicitly stated.
    3. Follow the provided schema exactly.
    4. Return only valid JSON matching the schema.
    5. If a field is not available in the job description, return null or an empty list as appropriate according to the schema.
    6. Preserve the original meaning of the job description.
    7. Extract technical skills, qualifications, experience requirements, responsibilities, and other relevant information accurately.
    8. Do not add any fields that are not defined in the schema.

    Schema:
    {schema}
    """

    jd_message = {
        "role": "user",
        "content": f"Job Description:\n\n{job_description}"
    }

    jd_messages = [
        {
            "role": "system",
            "content": jd_system_prompt
        },
        jd_message
    ]

    text = llm_call(jd_messages)
    raw = json.loads(text)
    return JobDes(**raw) # returns JobDes object, not a dictionary


def score_resume(job, resume):
    match_schema = MatchResult.model_json_schema()

    system_prompt = f"""
    You are an experienced HR resume screening engine.

    Evaluate the candidate resume against the job description and return JSON matching this schema:
    {json.dumps(match_schema, indent=2)}

    The "score" field must be a float from 0 to 100.
    The "details" field must be a JSON object containing:
    - candidate_name: string
    - matching_skills: list of strings
    - missing_skills: list of strings
    - experience_met: boolean
    - verdict: string (one concise sentence)

    Scoring guidance:
    90-100: Excellent match
    75-89:  Strong match
    60-74:  Moderate match
    40-59:  Weak match
    0-39:   Poor match

    Return valid JSON only. No markdown, no explanation.
    """

    user_prompt = f"""
    Job description:
    {json.dumps(job.model_dump(), indent=2)}

    Resume:
    {json.dumps(resume.model_dump(), indent=2)}
    """
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user",   "content": user_prompt},
    ]

    text = llm_call(messages)
    raw = json.loads(text)
    return MatchResult(**raw)



# print("Metadata:")
# print(result["metadata"])

# print(f"\nPages: {result['num_pages']}")
# print("\nFirst Page:")
# print(result["pages"][0]["text"][:500])
# print(f"\n Parsed Resume: {parsed_resume_text}")
# print(f"\n Parsed Job Desc: {parsed_job_des}")
# print(f"\n Score Resume: {resume_score}")

from pathlib import Path
folder = Path("resumes")

parsed_job_des = parse_job(job_description)
# print(parsed_job_des.minimum_experience)

all_results=[]
for file in folder.iterdir():

    # Process only PDF and DOCX files
    if file.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print(f"Processing: {file.name}")

    if file.suffix.lower() == ".pdf":
        parsed_text = parse_pdf(file)
    elif file.suffix.lower() == ".docx":
        parsed_text = parse_docx(file)
    else:
        continue


    parsed_resume_text = parse_resume(parsed_text)

    resume_score = score_resume(parsed_job_des, parsed_resume_text)
    
    print("Score:", resume_score.score)
    
    all_results.append({
        "name": parsed_resume_text.name,
        "score": resume_score.score,
        "details": resume_score.details
    })

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)

# Print final ranked list
print("\n── Candidate Rankings ──────────────────────────────")
for rank, candidate in enumerate(all_results, 1):
    print(f"  {rank}. {candidate['name']}  —  score: {candidate['score']}")
    print(f"     {candidate['details'].get('verdict', '')}")